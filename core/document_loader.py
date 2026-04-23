"""
Estrazione testo da file di vario formato.
Usato per preparare il contesto da inviare all'LLM.
"""

from pathlib import Path
from typing import Union


def load_document(file_path: Union[str, Path]) -> str:
    """
    Estrae il testo da un documento. Supporta pdf, docx, xlsx, csv, txt, md.
    Ritorna una stringa. Se il file è sconosciuto, prova a leggerlo come testo.
    """
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(file_path)

    ext = file_path.suffix.lower()

    if ext == ".pdf":
        return _load_pdf(file_path)
    if ext == ".docx":
        return _load_docx(file_path)
    if ext in (".xlsx", ".xlsm"):
        return _load_xlsx(file_path)
    if ext == ".csv":
        return _load_csv(file_path)
    if ext in (".txt", ".md"):
        return file_path.read_text(encoding="utf-8", errors="replace")

    # fallback: prova testuale
    try:
        return file_path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return f"[Impossibile leggere il file {file_path.name}: {e}]"


def _load_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        parts.append(f"--- Pagina {i} ---\n{text}")
    return "\n\n".join(parts)


def _load_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # includiamo anche le tabelle, rilevanti per estratti PR/PO
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def _load_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## Foglio: {sheet_name}")
        for row in ws.iter_rows(values_only=True):
            if any(v is not None for v in row):
                parts.append(
                    "\t".join("" if v is None else str(v) for v in row)
                )
    return "\n".join(parts)


def _load_csv(path: Path) -> str:
    import csv

    rows = []
    with path.open(newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for r in reader:
            rows.append("\t".join(r))
    return "\n".join(rows)
